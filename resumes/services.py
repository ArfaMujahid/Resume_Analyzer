import os
import re
import hashlib
from io import BytesIO
from typing import Tuple, Optional, Dict, Any

import PyPDF2
from docx import Document
from PIL import Image
import pytesseract

from django.conf import settings
from django.core.files.base import ContentFile
from .models import ResumeDocument, ParsedResume
from security.file_scanner import file_scanner


class DocumentParserService:
    """Service for parsing various document formats and extracting text"""

    SUPPORTED_TYPES = {
        'pdf': 'application/pdf',
        'doc': 'application/msword',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'txt': 'text/plain'
    }

    @staticmethod
    def get_file_hash(file_content: bytes) -> str:
        """Generate SHA-256 hash of file content"""
        return hashlib.sha256(file_content).hexdigest()

    @staticmethod
    def validate_file(file_content: bytes, filename: str) -> Tuple[bool, str]:
        """Validate file type and size"""
        import tempfile

        # Check file size (default 20MB)
        max_size = getattr(settings, 'MAX_FILE_SIZE', 20 * 1024 * 1024)
        if len(file_content) > max_size:
            return False, f"File size exceeds {max_size / (1024*1024):.0f}MB limit"

        # Check file extension
        file_ext = os.path.splitext(filename)[1].lower().lstrip('.')
        if file_ext not in DocumentParserService.SUPPORTED_TYPES:
            return False, f"Unsupported file type: {file_ext}"

        # Security scan the file
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_ext}') as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name

        try:
            is_safe, error_message, warnings = file_scanner.scan_file(temp_file_path, filename)

            if not is_safe:
                return False, f"Security check failed: {error_message}"

            # Log any warnings
            for warning in warnings:
                print(f"Security warning for {filename}: {warning}")

        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_file_path)
            except:
                pass

        return True, "File is valid"

    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> Tuple[str, str]:
        """Extract text from PDF content"""
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text = ""
            num_pages = len(pdf_reader.pages)

            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            # Clean up the text
            text = DocumentParserService._clean_text(text)

            # Check if extraction was successful
            if len(text.strip()) < 50:  # Very little text extracted
                return text, 'ocr_fallback'

            return text, 'text_extract'

        except Exception as e:
            print(f"PDF extraction error: {str(e)}")
            return "", 'failed'

    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> Tuple[str, str]:
        """Extract text from DOCX content"""
        try:
            doc = Document(BytesIO(file_content))

            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            # Clean up the text
            text = DocumentParserService._clean_text(text)

            return text, 'text_extract'

        except Exception as e:
            print(f"DOCX extraction error: {str(e)}")
            return "", 'failed'

    @staticmethod
    def extract_text_from_txt(file_content: bytes) -> Tuple[str, str]:
        """Extract text from plain text file"""
        try:
            # Try to decode as UTF-8 first
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                # Fallback to latin-1
                text = file_content.decode('latin-1')

            # Clean up the text
            text = DocumentParserService._clean_text(text)

            return text, 'text_extract'

        except Exception as e:
            print(f"TXT extraction error: {str(e)}")
            return "", 'failed'

    @staticmethod
    def ocr_fallback(file_content: bytes) -> Tuple[str, str]:
        """Use OCR to extract text from scanned documents"""
        try:
            # Convert PDF to images (first page only for now)
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            if len(pdf_reader.pages) == 0:
                return "", 'failed'

            # Get first page as image
            page = pdf_reader.pages[0]

            # This is a simplified approach - in production, you'd use pdf2image
            # For now, we'll return empty text as OCR requires additional setup
            text = ""

            # TODO: Implement proper OCR with pdf2image + pytesseract
            # Example:
            # from pdf2image import convert_from_bytes
            # images = convert_from_bytes(file_content)
            # for image in images:
            #     text += pytesseract.image_to_string(image) + "\n"

            return text, 'ocr_fallback'

        except Exception as e:
            print(f"OCR extraction error: {str(e)}")
            return "", 'failed'

    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""

        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Remove page numbers and headers
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)

        # Remove email addresses (optional, for privacy)
        # text = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '[EMAIL]', text)

        # Remove phone numbers (optional)
        # text = re.sub(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', '[PHONE]', text)

        # Fix common PDF extraction issues
        text = text.replace('-\n ', '')  # Remove hyphenation line breaks
        text = text.replace('\n ', '\n')  # Remove leading spaces after newlines

        return text.strip()

    @staticmethod
    def parse_document(resume_document: ResumeDocument) -> bool:
        """Main method to parse a document and save results"""
        try:
            # Update status to parsing
            resume_document.status = 'parsing'
            resume_document.save()

            # Read file content
            with open(resume_document.storage_ref, 'rb') as f:
                file_content = f.read()

            # Extract text based on file type
            file_ext = resume_document.file_type.lower()

            if file_ext == 'pdf':
                text, method = DocumentParserService.extract_text_from_pdf(file_content)
            elif file_ext in ['doc', 'docx']:
                text, method = DocumentParserService.extract_text_from_docx(file_content)
            elif file_ext == 'txt':
                text, method = DocumentParserService.extract_text_from_txt(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")

            # If text extraction failed, try OCR
            if not text or len(text.strip()) < 50:
                text, method = DocumentParserService.ocr_fallback(file_content)

            if not text:
                resume_document.status = 'failed'
                resume_document.parsing_error = "Could not extract text from document"
                resume_document.save()
                return False

            # Create or update parsed resume
            parsed_resume, created = ParsedResume.objects.get_or_create(
                resume=resume_document,
                defaults={
                    'extraction_method': method,
                    'raw_text': text[:10000],  # Store first 10k chars temporarily
                    'structured_json': {},
                    'section_index': {},
                    'skills_normalized': [],
                    'titles_normalized': [],
                    'companies': [],
                    'employment_history': [],
                    'education': [],
                    'certifications': [],
                    'quality_flags': {},
                    'embedding_refs': []
                }
            )

            if not created:
                parsed_resume.extraction_method = method
                parsed_resume.raw_text = text[:10000]
                parsed_resume.save()

            # Update status
            resume_document.status = 'parsed'
            resume_document.save()

            # Trigger next processing step (structuring)
            from .tasks import structure_resume_task
            structure_resume_task.delay(parsed_resume.id)

            return True

        except Exception as e:
            resume_document.status = 'failed'
            resume_document.parsing_error = str(e)
            resume_document.save()
            print(f"Document parsing error: {str(e)}")
            return False

    @staticmethod
    def estimate_extraction_quality(text: str) -> Dict[str, Any]:
        """Estimate the quality of text extraction"""
        if not text:
            return {'score': 0, 'issues': ['No text extracted']}

        issues = []
        score = 100

        # Check for common extraction issues
        if len(text) < 200:
            issues.append('Very little text extracted')
            score -= 50

        # Check for garbled text
        garbled_ratio = sum(1 for c in text if ord(c) < 32 or ord(c) > 126) / len(text)
        if garbled_ratio > 0.1:
            issues.append('High ratio of non-printable characters')
            score -= 30

        # Check for section headers
        section_keywords = ['experience', 'education', 'skills', 'summary', 'objective']
        has_sections = any(keyword.lower() in text.lower() for keyword in section_keywords)
        if not has_sections:
            issues.append('No clear section headers detected')
            score -= 20

        # Check for bullet points
        if '•' not in text and '-' not in text and '*' not in text:
            issues.append('No bullet points detected')
            score -= 10

        return {
            'score': max(0, score),
            'issues': issues,
            'text_length': len(text),
            'word_count': len(text.split())
        }